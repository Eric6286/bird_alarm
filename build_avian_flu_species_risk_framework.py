from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/wuyang/Documents/New project/广东省重要迁徙鸟类疫病监测预警体系与风险评估_修改稿.docx")

ACCENT = "111111"
LIGHT = "F2F2F2"
GRID = "C9C9C9"
DARK = "1F2933"
MUTED = "5F6B63"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def table_width(table, cm):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tbl_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def para(doc, text="", style=None, align=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, bold=True, color=ACCENT)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table_width(table, sum(widths))
    borders(table)
    repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_width(cell, widths[i])
        shade(cell, LIGHT)
        margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.5, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell_width(cell, widths[i])
            margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == 0 and len(val) <= 8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_font(r, size=9)
    para(doc, "", after=4)
    return table


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.color.rgb = RGBColor.from_string(DARK)
    header = section.header.paragraphs[0]
    header.text = "广东省重要迁徙鸟类疫病监测预警体系与风险评估"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = footer.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def build():
    doc = Document()
    configure(doc)

    title = para(doc, "广东省重要迁徙鸟类疫病监测预警体系与风险评估", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    set_run_font(title.runs[0], size=20, bold=True, color=ACCENT)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["基本思路", "先建立指标体系，再选择鸟类数量多、代表性强、资料可补充的地点，最后根据各地点鸟类名录和监测资料计算风险值。"],
            ["方法参照", "可参考鸟撞风险模型的做法，把一个地点内不同鸟种的数量、出现频次、季节和接触场景转化为可加总的风险贡献。"],
            ["资料口径", "年度水鸟调查、保护区调查、监测站记录和公开鸟类名录均可作为依据；单一年份调查只作为部分证据，不作为唯一依据。"],
            ["后续工作", "待补充三个代表地点的鸟类名录、优势物种、数量或频次、季节变化后，再计算地点风险值并进行比较。"],
        ],
        [3.0, 13.6],
    )

    heading(doc, "一、总体思路")
    para(doc, "本体系调整为“先建指标、再选地点、再计算风险值”的思路。前期不急于给出最终风险排序，而是先明确需要收集哪些数据、如何把鸟类名录转化为可计算变量，以及代表地点之间如何进行横向比较。")
    add_table(
        doc,
        ["步骤", "主要任务", "输出结果"],
        [
            ["1", "构建指标体系，明确鸟类、栖息地、接触场景、监测基础和异常信号等核心变量。", "全省通用的风险评价指标体系"],
            ["2", "选择代表性地点，优先考虑鸟类数量多、湿地类型典型、资料可获得、监测基础较清楚的区域。", "粤东、珠江口/大湾区、粤西代表样区"],
            ["3", "补充地点鸟类名录，记录每个地点的物种、类群、数量或出现频次、季节、优势物种等。", "地点—物种基础数据库"],
            ["4", "参考鸟撞风险模型，将单个物种的风险贡献加总为地点风险值。", "代表地点综合风险值"],
            ["5", "结合异常死亡、疑似或确诊结果等事件信号，对计算结果进行年度校正和预警触发。", "年度风险评估和预警建议"],
        ],
        [1.2, 10.2, 5.2],
    )

    heading(doc, "二、代表地点选择原则")
    para(doc, "代表地点不是由某一份年度调查报告单独推导出来，而是根据广东省沿海水鸟长期分布特点、重要湿地格局、鸟类数量、栖息地代表性和现有监测基础综合确定。年度调查结果可作为补充证据，用于说明这些地点具有较高鸟类数量或持续监测价值。")
    add_table(
        doc,
        ["区域", "代表地点", "选择依据", "后续需补充"],
        [
            ["粤东", "广东海丰湿地", "鸟类数量较多，滨海湿地和鱼塘复合生境典型，具有粤东水鸟停歇和越冬地代表性。", "完整鸟类名录、优势物种、数量等级、季节变化、异常记录"],
            ["珠江口 / 大湾区", "深圳湾内伶仃福田湿地", "鸟类数量不低，城市湿地、人鸟接触和公众活动场景明显，适合作为大湾区代表样区。", "鸟类名录、重点类群、公众活动场景、接触界面、救护记录"],
            ["粤西", "湛江雷州湾", "水鸟数量高，沿海滩涂和红树林湿地连续分布，具有粤西迁徙和越冬水鸟聚集代表性。", "鸟类名录、优势物种、迁徙停歇情况、年度数量变化、异常死亡记录"],
        ],
        [2.4, 3.2, 6.2, 4.8],
    )

    heading(doc, "三、风险值计算方法")
    para(doc, "可将每个代表地点看作一个样区，将样区内每个鸟种看作一个风险贡献单元。参考鸟撞风险模型中“物种风险贡献可加总”的处理方式，先计算单个物种的风险贡献值，再汇总为地点风险值。")
    para(doc, "建议采用简化公式：地点风险值 Rsite = Σ（Ai × Hi × Si × Ci） × M。")
    add_table(
        doc,
        ["变量", "含义", "建议取值方式", "说明"],
        [
            ["Ai", "物种数量或出现强度", "按数量等级、出现频次或优势度赋值，例如1—5分。", "解决“这个地点鸟多不多、该物种常不常见”的问题。"],
            ["Hi", "物种宿主关注系数", "按水鸟类群和公开资料中对疫病监测的关注程度赋值。", "不需要先写死具体病原学结论，可先按类群和专家意见给分。"],
            ["Si", "季节与聚集系数", "按迁徙季、越冬期、繁殖期或全年留居情况赋值。", "体现同一物种在不同季节对监测预警的意义不同。"],
            ["Ci", "接触场景系数", "按是否靠近公众活动区、养殖界面、鱼塘、水库、城市湿地等赋值。", "体现鸟群活动与人、家禽或管理场景的空间关系。"],
            ["M", "监测修正系数", "按监测站覆盖、数据连续性、异常死亡记录和检测结果反馈情况修正。", "用于把地点监测基础和事件信号纳入最终结果。"],
        ],
        [1.5, 3.4, 6.1, 5.6],
    )
    para(doc, "其中 Ai、Hi、Si、Ci 可先采用1—5分，M 可设为0.8—1.2的修正系数。资料不足时应标注“待补充”或采用保守默认值，避免为了计算而填入无法核实的数据。")

    heading(doc, "四、鸟类名录补充表")
    para(doc, "后续查询三个地点鸟类名录时，建议按下表整理。只要每个地点能形成同一格式的数据表，就可以比较不同地点的风险值。")
    add_table(
        doc,
        ["字段", "填写内容", "用途"],
        [
            ["地点", "广东海丰湿地、深圳湾内伶仃福田湿地、湛江雷州湾", "用于分组计算地点风险值"],
            ["中文名 / 学名", "鸟类物种名称，建议保留学名便于后续核对", "建立地点—物种清单"],
            ["类群", "雁鸭类、鸻鹬类、鸥类、鹭类、秧鸡类、其他水鸟等", "用于赋予宿主关注系数"],
            ["数量或频次", "实际数量、数量等级、出现频次或优势度", "用于计算 Ai"],
            ["季节", "越冬、迁徙停歇、夏候、留鸟或全年记录", "用于计算 Si"],
            ["主要生境", "滩涂、红树林、鱼塘、水库、河口、城市湿地等", "用于判断接触场景"],
            ["接触界面", "公众活动、观鸟、投喂、养殖邻近、救护记录等", "用于计算 Ci"],
            ["备注", "异常死亡、历史重点记录、资料来源、是否需核对", "用于年度校正和资料追溯"],
        ],
        [3.0, 8.0, 5.6],
    )

    heading(doc, "五、指标体系与权重建议")
    para(doc, "在已有地点名录和物种贡献值基础上，可再汇总为六类一级指标。权重不宜过细，先保证数据能查到、能解释、能年度更新。")
    add_table(
        doc,
        ["一级指标", "建议权重", "与名录计算的关系"],
        [
            ["鸟类数量与出现强度", "0.25", "主要由各物种 Ai 汇总形成，是地点鸟类数量多寡的核心指标。"],
            ["重点水鸟类群与物种组成", "0.20", "主要由 Hi 体现，依赖后续补充的鸟类名录和优势物种。"],
            ["季节聚集与迁徙使用", "0.15", "主要由 Si 体现，关注越冬期和迁徙停歇期的聚集情况。"],
            ["栖息地与接触场景", "0.15", "主要由 Ci 体现，关注城市湿地、公众活动、养殖邻近和共用水体等。"],
            ["监测站覆盖与数据质量", "0.10", "作为 M 的组成部分，反映资料连续性和可核查程度。"],
            ["异常死亡与检测结果", "0.15", "作为 M 的组成部分，同时保留直接触发专项评估的作用。"],
        ],
        [6.0, 2.2, 8.4],
    )

    heading(doc, "六、风险分级与触发机制")
    para(doc, "地点风险值可先用于三个代表地点之间的相对比较，待年度数据稳定后再设置固定阈值。初期建议采用相对分级：三个地点按风险值由高到低排序，结合异常事件和监测数据完整性确定巡查和资料补充优先级。")
    add_table(
        doc,
        ["结果类型", "判定方式", "建议应用"],
        [
            ["相对高值地点", "在三个代表地点中风险值最高，或某类物种贡献明显集中。", "优先补充名录、加密巡查、核对优势物种和接触场景。"],
            ["中等关注地点", "风险值居中，且无明显异常事件。", "维持常规监测，重点补齐缺失字段。"],
            ["相对低值地点", "风险值较低，且鸟类数量、接触场景或异常记录均较少。", "保留常规监测，不作为年度重点加密对象。"],
            ["直接触发事件", "出现疑似或确诊高致病性禽流感、聚集性死亡事件或连续异常死亡记录。", "不受相对排序限制，直接进入专项应急响应评估或会商。"],
        ],
        [3.0, 6.0, 7.6],
    )

    heading(doc, "七、年度更新")
    para(doc, "本体系的关键不是一次性算出最终答案，而是形成可持续更新的数据表。每年可根据三个代表地点新增鸟类名录、数量变化、优势物种、异常死亡记录和检测结果，更新物种贡献值和地点风险值。若后续发现其他地点鸟类数量更高或资料更完整，也可按同一方法替换或新增代表地点。")
    para(doc, "本体系主要用于监测预警和风险研判，不替代国家和广东省现行动物疫病、野生动物疫源疫病及突发公共卫生事件相关法定报告和应急处置程序；涉及疑似或确诊疫情的，应按现行规定执行。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

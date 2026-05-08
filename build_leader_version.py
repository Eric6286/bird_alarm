from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/wuyang/Documents/New project/广东省重要迁徙鸟类疫病监测预警体系与风险评估_领导版修改稿.docx")

LIGHT = "F2F2F2"
GRID = "C9C9C9"
DARK = "1F2933"
MUTED = "666666"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_cm):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def paragraph(doc, text="", style=None, align=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, bold=True, color="111111")
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, sum(widths))
    set_borders(table)
    repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.5, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == 0 and len(str(text)) <= 8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, size=9)
    paragraph(doc, "", after=3)
    return table


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.color.rgb = RGBColor.from_string(DARK)
    header = section.header.paragraphs[0]
    header.text = "广东省重要迁徙鸟类疫病监测预警体系与风险评估"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = footer.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def build():
    doc = Document()
    configure(doc)

    title = paragraph(doc, "广东省重要迁徙鸟类疫病监测预警体系与风险评估", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    set_run_font(title.runs[0], size=20, bold=True, color="111111")

    heading(doc, "一、总体框架")
    paragraph(doc, "本体系面向广东省重要迁徙鸟类疫病监测预警和风险评估工作，围绕鸟类迁徙路线、主要分布物种、重点栖息地、监测站布局和异常事件记录，建立可用于重点区域识别、重点关注对象筛选和风险指数计算的技术指标体系。")
    paragraph(doc, "体系重点服务于三个方面：一是识别广东省内迁徙水鸟集中分布和停歇越冬的重要区域；二是根据鸟类名录和主要类群提出重点关注对象；三是参考风电鸟撞风险评价中按物种贡献累加的思路，计算代表地点的禽流感风险指数，为监测布点和年度预警研判提供依据。")
    add_table(
        doc,
        ["工作环节", "主要内容", "形成结果"],
        [
            ["基础资料整理", "汇集迁徙路线、广东水鸟分布、监测站名录、越冬水鸟调查、异常死亡和检测结果等资料。", "广东省迁徙鸟类与监测基础资料库"],
            ["指标体系构建", "围绕鸟类数量、物种组成、季节聚集、栖息地接触、监测基础和异常事件建立指标。", "风险评估主要技术指标体系"],
            ["重点对象识别", "根据水鸟类群、物种名录、数量和出现频次筛选重点关注类群和物种。", "重点关注对象清单"],
            ["风险指数计算", "对代表地点内各物种风险贡献进行赋值和汇总，形成地点风险指数。", "代表地点风险评估结果"],
        ],
        [3.0, 8.4, 5.2],
    )

    heading(doc, "二、广东省迁徙鸟类与疫病风险基础")
    paragraph(doc, "广东位于东亚—澳大利西亚候鸟迁徙路线上，沿海滩涂、河口、红树林、鱼塘、水库和城市湿地为迁徙水鸟提供停歇、觅食和越冬场所。公开监测资料显示，广东已记录野生鸟类584种，其中迁徙鸟类412种；近年越冬水鸟同步监测覆盖沿海和内陆多类湿地，持续补充水鸟种类、数量、分布和栖息地变化等基础数据。")
    paragraph(doc, "从疫病监测预警角度看，广东的风险评估不宜只从行政管理能力或一般处置流程出发，而应把鸟类迁徙路线、重点湿地、主要水鸟类群和现有监测站布局作为基础。对水鸟数量较多、类群组成较复杂、与公众活动或养殖水体存在接触场景的区域，应优先纳入年度风险评估。")
    add_table(
        doc,
        ["资料类别", "广东相关内容", "在风险评估中的用途"],
        [
            ["迁徙路线", "东亚—澳大利西亚候鸟迁徙路线经过广东沿海和内陆湿地。", "确定重点区域和季节性监测窗口"],
            ["水鸟分布", "沿海滩涂、河口、红树林、鱼塘、水库和城市湿地是越冬和迁徙水鸟集中区域。", "判断水鸟聚集强度和潜在接触场景"],
            ["主要类群", "重点关注雁鸭类、鸻鹬类、鸥类等水鸟类群。", "形成后续物种名录赋值基础"],
            ["监测基础", "全省已建国家级和省级陆生野生动物疫源疫病监测站，部分重点湿地具有连续鸟类监测资料。", "提高风险指数计算的可核查性"],
        ],
        [3.0, 7.2, 6.4],
    )

    heading(doc, "三、风险评估指标体系")
    paragraph(doc, "指标设置应突出可获得、可更新、可计算。对于基层难以长期稳定获取的数据，不作为核心判别要点；对于鸟类名录、数量或频次、季节变化、栖息地类型、监测站覆盖和异常事件等可通过调查、监测和资料整理获得的数据，作为核心指标。")
    add_table(
        doc,
        ["一级指标", "二级指标", "主要数据来源", "评分说明"],
        [
            ["鸟类数量与出现频率", "水鸟总量、重点类群数量、年度或季节出现频次", "越冬水鸟同步监测、保护区调查、固定样线或样点记录", "数量越高、出现越稳定，分值越高"],
            ["重点类群与物种组成", "雁鸭类、鸻鹬类、鸥类及其他重点水鸟物种组成", "地点鸟类名录、保护区名录、年度调查记录", "重点类群占比高或重点物种较多，分值越高"],
            ["迁徙与越冬季节聚集", "迁徙停歇期、越冬期、短期聚集高峰", "月度巡护、同步监测、历史调查资料", "迁徙或越冬高峰明显，分值越高"],
            ["栖息地与接触场景", "滩涂、红树林、鱼塘、水库、城市湿地，以及公众活动、养殖邻近、共用水体等", "湿地资源资料、现场踏查、保护区巡护、属地资料", "接触场景越明显，分值越高"],
            ["监测站覆盖与资料质量", "监测站覆盖、数据连续性、物种名录完整性、异常记录完整性", "监测站名录、保护区台账、年度报告", "监测基础越清楚，结果可信度越高"],
            ["异常死亡与检测结果", "病弱或死亡野鸟记录、聚集性死亡事件、疑似或确诊结果", "巡护记录、公众报告、救护机构记录、实验室报告", "作为风险升级和专项评估的重要触发信号"],
        ],
        [3.2, 4.1, 4.8, 4.5],
    )

    heading(doc, "四、重点关注对象")
    paragraph(doc, "重点关注对象应根据广东迁徙水鸟分布、代表地点鸟类名录和年度监测结果确定。当前阶段可先从类群层面提出重点对象，后续在海丰湿地、深圳湾内伶仃福田湿地、湛江雷州湾等代表地点补充完整鸟类名录后，再细化到具体物种。")
    add_table(
        doc,
        ["重点类群", "关注依据", "后续名录整理要点"],
        [
            ["雁鸭类", "与水库、鱼塘、河口、开阔水面等生境关系密切，迁徙和越冬季节数量变化明显。", "记录物种名、数量或频次、越冬或迁徙停歇情况、是否形成集中活动区域"],
            ["鸻鹬类", "多利用河口、滩涂和潮间带湿地，部分沿海区域可形成较高数量聚集。", "记录滩涂利用情况、优势物种、潮汐或季节变化、与公众活动或养殖区距离"],
            ["鸥类", "常在河口、近岸滩涂、水面和人类活动较多区域出现，部分区域数量较高。", "记录数量峰值、活动水域、取食场景、与码头、渔港或城市湿地的关系"],
            ["其他重点水鸟", "包括在代表地点数量较高、保护等级较高、或在异常事件中需要重点关注的水鸟物种。", "根据各地点名录和年度监测结果动态补充，不预先扩大到难以核实的所有物种"],
        ],
        [3.0, 6.2, 7.4],
    )

    heading(doc, "五、风险指数计算方法")
    paragraph(doc, "参考风电鸟撞风险评价中按物种风险贡献累加的思路，可将每个代表地点视为一个评价单元，将地点鸟类名录中的每个物种视为一个风险贡献单元。地点风险指数由各物种风险贡献值汇总形成，并根据监测基础和异常事件进行修正。")
    paragraph(doc, "建议计算公式为：R = Σ（Ai × Bi × Ci × Di） × M。")
    add_table(
        doc,
        ["变量", "含义", "建议赋值"],
        [
            ["Ai", "物种数量或出现频率", "按实测数量、数量等级、出现频次或优势度赋1—5分"],
            ["Bi", "重点类群或物种关注系数", "雁鸭类、鸻鹬类、鸥类及重点物种按关注程度赋1—5分"],
            ["Ci", "季节聚集系数", "越冬期、迁徙停歇期或短期高峰明显时赋较高值"],
            ["Di", "栖息地与接触场景系数", "与公众活动、养殖界面、鱼塘、水库、河口滩涂等接触场景越明显，赋值越高"],
            ["M", "监测与异常事件修正系数", "监测资料完整且无异常可取1.0；出现异常死亡、疑似或确诊结果时上调"],
        ],
        [1.5, 6.4, 8.7],
    )
    paragraph(doc, "在实际计算中，可先按物种逐项赋值，再按地点汇总。若某一地点暂未取得完整物种名录，可先完成资料表结构和已有水鸟数量记录，待名录补齐后再计算正式风险指数。")

    heading(doc, "六、资料收集与年度更新")
    paragraph(doc, "为避免指标体系停留在概念层面，各代表地点应按统一格式整理鸟类名录和监测资料。资料收集不宜追求过多难以获得的管理性指标，应优先保障鸟类、栖息地和监测结果三类数据的完整性。")
    add_table(
        doc,
        ["字段", "填写内容", "用途"],
        [
            ["地点", "海丰湿地、深圳湾内伶仃福田湿地、湛江雷州湾等", "分地点计算风险指数"],
            ["物种名称", "中文名、学名、保护等级或重点关注说明", "建立物种清单"],
            ["类群", "雁鸭类、鸻鹬类、鸥类、其他重点水鸟", "确定重点类群系数"],
            ["数量或频次", "实测数量、数量等级、出现频次、优势度", "计算Ai"],
            ["季节", "越冬、迁徙停歇、夏候、留鸟或全年记录", "计算Ci"],
            ["生境和接触场景", "滩涂、红树林、鱼塘、水库、城市湿地、公众活动、养殖邻近等", "计算Di"],
            ["监测和异常记录", "监测站覆盖、异常死亡、疑似或确诊结果、资料来源", "确定M和触发专项评估"],
        ],
        [3.0, 8.2, 5.4],
    )

    heading(doc, "七、代表性重点区域及示例计算")
    paragraph(doc, "在指标体系和计算方法明确后，可选择鸟类数量较多、区域代表性较强、监测基础较好的地点开展示例计算。根据广东沿海水鸟长期分布特点、现有监测站布局及年度水鸟调查资料，建议先选取粤东广东海丰湿地、珠江口/大湾区深圳湾内伶仃福田湿地、粤西湛江雷州湾作为代表性重点区域。")
    add_table(
        doc,
        ["区域", "代表地点", "已有监测基础", "选择理由", "名录补充重点"],
        [
            ["粤东", "广东海丰湿地", "广东海丰保护区国家级陆生野生动物疫源疫病监测站；已有越冬水鸟调查记录约7930只，历史记录数量较高。", "滨海湿地和鱼塘复合生境典型，鸟类数量较多，适合作为粤东代表点。", "雁鸭类、鸻鹬类、鸥类及优势水鸟物种"],
            ["珠江口 / 大湾区", "深圳湾内伶仃福田湿地", "广东内伶仃岛福田保护区国家级陆生野生动物疫源疫病监测站；已有越冬水鸟调查记录约5363只。", "城市湿地特征明显，鸟类数量不低，人鸟接触和公众活动场景突出。", "重点水鸟名录、数量频次、公众活动和接触场景"],
            ["粤西", "湛江雷州湾", "广东湛江保护区国家级陆生野生动物疫源疫病监测站；已有越冬水鸟调查记录约15929只。", "沿海滩涂和红树林湿地连续分布，水鸟数量高，粤西代表性强。", "鸻鹬类、鸥类、雁鸭类及迁徙停歇物种"],
        ],
        [2.0, 2.8, 4.6, 4.0, 3.2],
    )
    paragraph(doc, "示例计算时，可先对每个代表地点建立“地点—物种”表，再按物种逐项赋值。以下为计算表样式，待具体鸟类名录补充后形成正式风险指数。")
    add_table(
        doc,
        ["地点", "物种或类群", "Ai数量/频次", "Bi关注系数", "Ci季节系数", "Di接触系数", "物种贡献值"],
        [
            ["海丰湿地", "待补充具体物种", "按数量或频次赋值", "按类群或物种赋值", "按越冬/迁徙情况赋值", "按生境和接触场景赋值", "Ai×Bi×Ci×Di"],
            ["深圳湾内伶仃福田湿地", "待补充具体物种", "按数量或频次赋值", "按类群或物种赋值", "按越冬/迁徙情况赋值", "按生境和接触场景赋值", "Ai×Bi×Ci×Di"],
            ["湛江雷州湾", "待补充具体物种", "按数量或频次赋值", "按类群或物种赋值", "按越冬/迁徙情况赋值", "按生境和接触场景赋值", "Ai×Bi×Ci×Di"],
        ],
        [2.8, 3.0, 2.3, 2.1, 2.1, 2.3, 2.0],
    )
    paragraph(doc, "地点风险指数按各物种贡献值汇总后乘以监测与异常事件修正系数M。三个代表地点的风险值可用于年度横向比较，优先识别需要补充物种名录、加密巡查或开展专项会商的区域。")

    heading(doc, "八、预警应用与运行")
    paragraph(doc, "风险指数主要用于监测预警和年度布点，不替代国家和广东省现行动物疫病、野生动物疫源疫病及突发公共卫生事件相关法定报告和应急处置程序。疑似或确诊高致病性禽流感、聚集性死亡事件、连续异常死亡记录等，应直接进入专项应急响应评估或会商。")
    paragraph(doc, "年度运行中，应在迁徙季前更新代表地点鸟类名录和监测点位，迁徙和越冬高峰期开展滚动评估，年度结束后复核指标权重、重点关注对象和代表地点适用性。")

    heading(doc, "九、主要资料依据")
    add_table(
        doc,
        ["资料", "用途"],
        [
            ["广东省林业科技创新项目合同书", "明确项目需提出监测预警和风险评估主要技术指标体系。"],
            ["广东省陆生野生动物疫源疫病监测站名录", "匹配代表地点与现有监测站基础。"],
            ["广东省越冬水鸟同步监测资料", "提供水鸟数量、监测单元和重点区域参考。"],
            ["国家林草局、广东省林业局等公开资料", "支撑广东迁徙通道、野生鸟类和越冬水鸟基础背景。"],
        ],
        [6.0, 10.6],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

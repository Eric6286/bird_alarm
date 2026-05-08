from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/wuyang/Documents/New project/广东省重要迁徙鸟类疫病监测预警体系与风险评估_鸟类名录待补充框架稿.docx")


ACCENT = "111111"
LIGHT = "F2F2F2"
GRID = "C9C9C9"
DARK = "1F2933"
MUTED = "5F6B63"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_cm):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, align=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, bold=True, color=ACCENT if level == 1 else DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, sum(widths))
    set_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == 0 and len(text) <= 8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=9)
    add_para(doc, "", space_after=4)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)

    styles = doc.styles
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "广东省重要迁徙鸟类疫病监测预警体系与风险评估"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = footer.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def build():
    doc = Document()
    configure_doc(doc)

    title = add_para(doc, "广东省重要迁徙鸟类疫病监测预警体系与风险评估", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    title.runs[0].font.size = Pt(20)
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = RGBColor.from_string(ACCENT)
    subtitle = add_para(doc, "重点区域与鸟类数据支撑框架（修改稿）", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["适用对象", "广东省迁徙水鸟及其主要停歇、越冬、觅食和人鸟接触区域的疫病监测预警。"],
            ["修改重点", "将原指标体系调整为以鸟类数据、重点湿地、监测站布局和异常事件为核心的可更新框架。"],
            ["资料基础", "广东省沿海水鸟长期分布特点、重要湿地格局、现有陆生野生动物疫源疫病监测站布局，以及年度水鸟监测资料。"],
            ["使用方式", "先以代表性重点区域建立资料框架，后续补充物种名录、优势物种、数量变化和年度检测结果后开展评分。"],
        ],
        [3.0, 13.6],
    )

    add_heading(doc, "一、总体框架", 1)
    add_para(
        doc,
        "本体系用于广东省重要迁徙鸟类疫病监测预警和风险研判。修改后的指标体系不另行构建新的地点风险指数，而是在原有预警框架中强化鸟类数据、广东沿海湿地格局和监测站可获得资料的支撑作用。体系重点关注迁徙路线和越冬地分布、重点水鸟栖息地、主要水鸟类群和重点宿主、现有疫源疫病监测站布局、异常死亡记录和检测结果等信息。",
    )
    add_para(
        doc,
        "综合广东省沿海水鸟分布特点、重要湿地格局、现有监测站布局及年度越冬水鸟同步监测结果，可先选取粤东海丰湿地、珠江口/大湾区深圳湾内伶仃福田湿地、粤西湛江雷州湾作为代表性重点区域。三个区域用于组织资料补充和指标评分，不表示其他区域不纳入常规监测。",
    )
    add_table(
        doc,
        ["评估维度", "一级指标", "主要数据基础", "应用场景"],
        [
            ["鸟类迁徙与区域重要性", "迁徙路线与区域重要性；水鸟聚集强度；重点宿主类群与物种组成", "迁徙通道资料、保护区鸟类调查、年度水鸟同步监测、固定样线或样点计数", "识别重点区域、安排巡查频次、确定年度补充名录"],
            ["栖息地与接触场景", "栖息地类型与人鸟接触场景", "湿地类型、滩涂/鱼塘/水库分布、公众活动、养殖和市场周边情况", "识别需要提示、隔离、巡护或现场管理的接触界面"],
            ["监测基础与事件信号", "监测站覆盖与数据可获得性；异常死亡野鸟与检测结果", "疫源疫病监测站名录、巡护台账、异常死亡记录、检测结果反馈", "支撑评分更新、会商研判和专项应急响应评估"],
        ],
        [3.6, 4.6, 4.6, 3.8],
    )

    add_heading(doc, "二、广东省代表性重点区域划定", 1)
    add_para(
        doc,
        "代表区域按“区域覆盖、鸟类数量、湿地代表性、监测基础和后续资料可补充性”综合选取。后续工作应围绕每个区域补充鸟类名录、优势物种、候鸟数量、季节变化和异常记录，使指标评分有据可查。",
    )
    add_table(
        doc,
        ["区域", "代表性重点区域", "对应监测基础", "后续需补充资料", "风险关注点"],
        [
            ["粤东", "广东海丰湿地", "广东海丰保护区国家级陆生野生动物疫源疫病监测站；海丰鸟类省级自然保护区长期巡护基础", "物种名录、优势水鸟、迁徙和越冬数量、季节变化、历史异常记录", "滨海湿地和鱼塘复合生境，多类群水鸟聚集"],
            ["珠江口 / 大湾区", "深圳湾内伶仃福田湿地", "广东内伶仃岛福田保护区国家级陆生野生动物疫源疫病监测站；深圳湾长期鸟类监测基础", "物种名录、重点宿主类群、人鸟接触场景、公众活动强度、救护记录", "城市湿地、人鸟接触明显，公众活动和管理界面复杂"],
            ["粤西", "湛江雷州湾", "广东湛江保护区国家级陆生野生动物疫源疫病监测站；湛江红树林保护区及沿海滩涂监测基础", "物种名录、优势物种、迁徙停歇情况、年度数量变化、异常死亡记录", "高数量水鸟聚集，沿海滩涂和红树林湿地连续分布"],
        ],
        [2.1, 3.2, 4.0, 3.9, 3.4],
    )

    add_heading(doc, "三、指标体系详细内容", 1)
    add_para(
        doc,
        "修改后建议保留六个一级指标。每个二级指标采用1—5分描述风险或关注程度，评分依据应尽量来自保护区巡护、监测站月报、年度水鸟调查、鸟类名录、异常死亡记录和检测结果等可获得资料。部门协同、培训和宣传不再作为一级指标，转入预警响应和运行保障部分。",
    )
    add_table(
        doc,
        ["一级指标", "二级指标", "观测或判定要点", "主要数据来源", "评分关注"],
        [
            ["1. 迁徙路线与区域重要性", "1.1 迁徙通道位置", "是否位于东亚—澳大利西亚迁飞通道及广东沿海重要停歇、越冬或补给区域", "迁徙通道资料、环志和卫星追踪资料、保护区记录", "通道位置越关键，关注分值越高"],
            ["1. 迁徙路线与区域重要性", "1.2 区域代表性", "是否能代表粤东、珠江口/大湾区或粤西等广东重点水鸟分布单元", "省内湿地格局、监测站名录、历史鸟类调查", "代表性越强，越适合作为重点样区"],
            ["1. 迁徙路线与区域重要性", "1.3 与周边区域连通性", "与河口、滩涂、鱼塘、水库、红树林等相邻生境是否形成连续活动空间", "遥感底图、保护区边界、巡护记录", "连通性强时应关注跨点位移动"],
            ["2. 水鸟聚集强度", "2.1 年度数量规模", "年度监测中水鸟总量及重点类群数量水平", "年度水鸟同步监测、固定样线/样点计数", "数量越高，现场巡查优先级越高"],
            ["2. 水鸟聚集强度", "2.2 季节变化", "秋冬越冬期、春季北迁期及异常天气后的数量变化", "月度巡护、保护区记录、气象资料", "短期快速增加时提高关注"],
            ["2. 水鸟聚集强度", "2.3 高密度聚集点", "是否形成稳定聚集的滩涂、鱼塘、河口、水库或湿地公园点位", "巡护台账、点位照片、无人机或望远镜观察记录", "聚集点稳定且密度高时提高分值"],
            ["3. 重点宿主类群与物种组成", "3.1 重点水鸟类群", "雁鸭类、鸻鹬类、鸥类、鹭类等水鸟类群的出现和占比", "鸟类名录、样线调查、保护区年度报告", "重点类群丰富或占比高时提高关注"],
            ["3. 重点宿主类群与物种组成", "3.2 优势物种", "代表区域内优势物种及其季节变化是否清楚", "后续补充物种名录、历史调查、观鸟记录辅助核对", "优势物种明确后可作为年度跟踪对象"],
            ["3. 重点宿主类群与物种组成", "3.3 异常个体记录", "是否记录病弱、行动异常、死亡或聚集异常的野鸟", "巡护员、公众报告、救护机构、监测站日报", "异常个体出现时触发核查"],
            ["4. 栖息地类型与人鸟接触场景", "4.1 栖息地类型", "红树林、河口滩涂、鱼塘、水库、湿地公园等生境类型及其组合", "保护区资料、湿地资源调查、现场踏查", "复合生境和高频利用生境需重点记录"],
            ["4. 栖息地类型与人鸟接触场景", "4.2 人鸟接触场景", "观鸟、游憩、投喂、捕鱼、养殖、市场周边等接触或扰动场景", "现场巡护、属地管理记录、公众活动信息", "接触越频繁，提示和管理需求越高"],
            ["4. 栖息地类型与人鸟接触场景", "4.3 周边养殖和共用水体", "水鸟活动区域与水禽养殖、鱼塘、沟渠或生活水体是否相邻", "农业农村部门台账、现场核查、水系图", "相邻或共用水体时提高关注"],
            ["5. 监测站覆盖与数据可获得性", "5.1 监测站覆盖", "代表区域及周边是否有国家级或省级陆生野生动物疫源疫病监测站支撑", "监测站名录、保护区管理机构资料", "覆盖清楚、责任主体明确时便于年度更新"],
            ["5. 监测站覆盖与数据可获得性", "5.2 数据连续性", "是否具备月度巡护、年度水鸟调查、异常事件记录等连续资料", "监测站月报、保护区台账、年度报告", "连续性越好，评分可信度越高"],
            ["5. 监测站覆盖与数据可获得性", "5.3 数据补充缺口", "物种名录、优势物种、数量变化或接触场景是否仍需补充", "资料清单、专家核对、现场补充调查", "缺口越大，先标注待补充而非强行评分"],
            ["6. 异常死亡野鸟与检测结果", "6.1 异常死亡记录", "同地点多只死亡、连续多日发现或伴随病弱个体等事件信号", "巡护记录、公众报告、救护机构记录", "聚集性死亡或连续发现应即时会商"],
            ["6. 异常死亡野鸟与检测结果", "6.2 报告和处置时效", "发现后是否完成定位、拍照、属地报告、临时提示和规范处置", "监测站台账、属地报告、现场记录", "时效越差，运行风险越高"],
            ["6. 异常死亡野鸟与检测结果", "6.3 检测结果反馈", "是否有疑似、确诊或其他异常检测结果，以及反馈至属地和相关部门的时效", "实验室报告、信息平台、会商记录", "疑似或确诊高致病性禽流感、聚集性死亡事件触发专项应急响应评估"],
        ],
        [3.0, 3.2, 4.0, 3.4, 3.0],
    )

    add_heading(doc, "四、评分与预警", 1)
    add_para(
        doc,
        "评分采用“二级指标1—5分、一级指标汇总、综合指数研判、异常事件直接触发”的方式。常态期可按季度更新，迁徙高峰期或出现异常事件时可按月或即时更新。资料不足的二级指标应标注“待补充”，不宜用无法核实的数据替代。",
    )
    add_bullet(doc, "二级指标评分：1分表示关注程度较低或资料显示风险较低，3分表示需持续关注，5分表示需要优先核查、加密巡查或组织会商。")
    add_bullet(doc, "一级指标得分：原则上取该一级指标下二级指标平均值；若出现4分及以上的异常信号，可按最高二级指标得分计入，避免被平均值稀释。")
    add_bullet(doc, "综合风险指数：R = Σ(Wi × Si)。Wi为一级指标权重，Si为一级指标得分，权重可根据年度复盘和资料完整度调整。")
    add_bullet(doc, "直接触发机制：疑似或确诊高致病性禽流感、聚集性死亡事件、连续异常死亡记录等，应直接进入专项应急响应评估或省级会商。")
    add_table(
        doc,
        ["一级指标", "建议权重", "权重设置理由"],
        [
            ["迁徙路线与区域重要性", "0.15", "体现广东沿海迁徙和越冬格局，是划定重点区域的基础。"],
            ["水鸟聚集强度", "0.20", "直接反映现场鸟群规模和巡查优先级。"],
            ["重点宿主类群与物种组成", "0.20", "体现鸟类数据支撑，是后续补充名录和优势物种的核心。"],
            ["栖息地类型与人鸟接触场景", "0.15", "用于识别人鸟接触、养殖邻近和公众活动管理需求。"],
            ["监测站覆盖与数据可获得性", "0.15", "保证评分有资料来源，避免指标体系停留在概念层面。"],
            ["异常死亡野鸟与检测结果", "0.15", "作为事件信号和预警升级的重要依据，同时保留直接触发机制。"],
        ],
        [6.2, 2.2, 8.2],
    )
    add_table(
        doc,
        ["综合风险指数R", "等级", "监测预警建议", "响应重点"],
        [
            ["1.0≤R<2.0", "常规关注", "维持常规巡护，更新重点区域基础资料和联系人清单。", "做好年度鸟类名录和监测数据归档。"],
            ["2.0≤R<3.0", "加强关注", "迁徙季适当加密重点湿地巡查，补齐优势物种、数量变化和接触场景资料。", "开展部门信息共享和重点点位提示。"],
            ["3.0≤R<4.0", "较高关注", "对代表区域及周边接触界面加密巡查，组织风险会商。", "设置现场提示，关注异常个体和周边养殖/公众活动。"],
            ["R≥4.0或出现直接触发事件", "专项评估", "即时核查异常死亡、疑似或确诊结果、聚集性死亡等事件。", "按现行规定启动专项应急响应评估和属地协同处置。"],
        ],
        [3.2, 2.1, 5.4, 5.9],
    )

    add_heading(doc, "五、资料补充和年度更新", 1)
    add_para(
        doc,
        "本体系为动态更新框架。后续可根据三个代表性重点区域的鸟类名录、年度水鸟数量、优势物种、重点宿主类群、异常死亡记录和检测结果，对指标评分和重点关注对象进行年度修订。建议每个代表区域建立一张可持续更新的资料表，先补齐“有什么鸟、数量如何、何时聚集、在哪里接触、由谁监测、是否有异常”六类信息。",
    )
    add_table(
        doc,
        ["代表区域", "需补充的鸟类资料", "需补充的监测资料", "年度更新产出"],
        [
            ["广东海丰湿地", "完整物种名录、优势水鸟、迁徙和越冬数量、重点类群季节变化", "海丰保护区巡护台账、异常死亡记录、检测结果摘要", "粤东重点区域年度鸟类与预警评分表"],
            ["深圳湾内伶仃福田湿地", "深圳湾及内伶仃福田湿地物种名录、优势物种、人鸟接触场景记录", "保护区和城市湿地巡护、公众活动、救护和异常报告记录", "珠江口/大湾区城市湿地年度风险关注清单"],
            ["湛江雷州湾", "雷州湾及周边红树林、滩涂水鸟名录、数量峰值、迁徙停歇情况", "湛江红树林保护区及相关监测站巡护、异常死亡、检测结果摘要", "粤西沿海滩涂年度鸟类与预警评分表"],
        ],
        [3.3, 4.7, 4.4, 4.2],
    )

    add_heading(doc, "六、指标体系运行", 1)
    add_para(doc, "（一）日常监测。各监测站围绕代表性重点区域及周边湿地、河口、滩涂、鱼塘、水库和城市湿地开展巡护记录，重点记录鸟类种类、数量、聚集点、异常个体和人鸟接触场景。")
    add_para(doc, "（二）资料汇集。省级或属地平台按月汇总监测站、保护区、救护机构、年度水鸟调查和检测结果摘要，形成可追溯的数据台账。")
    add_para(doc, "（三）评分会商。常态期按季度评分，迁徙高峰期按月评分；出现异常死亡、疑似或确诊结果、聚集性死亡事件时即时会商。")
    add_para(doc, "（四）分级响应。预警等级确定后，应同步明确责任单位、时间要求和现场措施。公众提示、人员培训和部门协同作为响应保障措施，不再单独作为一级评分指标。")
    add_para(doc, "（五）年度复盘。每年迁徙季结束后，复核代表区域、优势物种、数据缺口和权重设置，必要时调整重点区域清单和指标评分说明。")
    add_para(
        doc,
        "本指标体系主要用于监测预警和风险研判，不替代国家和广东省现行动物疫病、野生动物疫源疫病及突发公共卫生事件相关法定报告和应急处置程序；涉及疑似或确诊疫情的，应按现行规定执行。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
